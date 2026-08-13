#!/usr/bin/env python3
"""Build the compact IJCCE Sections 1-3 working manuscript from audited Markdown."""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
SOURCE = HERE / "IJCCE_Manuscript_Sections_1_3_Source.md"
OUTPUT = HERE / "IJCCE_Manuscript_Sections_1_3_Working_Draft_v6_Base.docx"


MATH_MARKER_RE = re.compile(r"\[\[MATH:([a-z0-9_]+)\]\]")
DISPLAY_MARKER_RE = re.compile(r"\[\[EQ:([a-z0-9_]+)\]\]")


def m_element(tag: str, *, value: Optional[str] = None):
    element = OxmlElement(f"m:{tag}")
    if value is not None:
        element.set(qn("m:val"), value)
    return element


def m_run(text: str, *, upright: bool = False):
    run = m_element("r")
    run_pr = m_element("rPr")
    style = m_element("sty", value="p" if upright else "i")
    run_pr.append(style)
    run.append(run_pr)
    node = m_element("t")
    node.text = text
    run.append(node)
    return run


def m_sequence(*parts):
    container = m_element("e")
    for part in parts:
        if part is None:
            continue
        if isinstance(part, str):
            part = m_run(part)
        append_math(container, part)
    return container


def append_math(container, expression):
    if expression.tag == qn("m:e"):
        for child in expression:
            container.append(deepcopy(child))
    else:
        container.append(deepcopy(expression))


def m_sub(base, subscript):
    node = m_element("sSub")
    base_node = m_element("e")
    append_math(base_node, base)
    sub_node = m_element("sub")
    append_math(sub_node, subscript)
    node.extend((base_node, sub_node))
    return node


def m_sup(base, superscript):
    node = m_element("sSup")
    base_node = m_element("e")
    append_math(base_node, base)
    sup_node = m_element("sup")
    append_math(sup_node, superscript)
    node.extend((base_node, sup_node))
    return node


def m_subsup(base, subscript, superscript):
    node = m_element("sSubSup")
    base_node = m_element("e")
    append_math(base_node, base)
    sub_node = m_element("sub")
    append_math(sub_node, subscript)
    sup_node = m_element("sup")
    append_math(sup_node, superscript)
    node.extend((base_node, sub_node, sup_node))
    return node


def m_fraction(numerator, denominator):
    node = m_element("f")
    props = m_element("fPr")
    props.append(m_element("type", value="bar"))
    num = m_element("num")
    append_math(num, numerator)
    den = m_element("den")
    append_math(den, denominator)
    node.extend((props, num, den))
    return node


def m_delimiter(expression, opening="(", closing=")"):
    node = m_element("d")
    props = m_element("dPr")
    props.extend((m_element("begChr", value=opening), m_element("endChr", value=closing)))
    expr = m_element("e")
    append_math(expr, expression)
    node.extend((props, expr))
    return node


def m_nary(expression, *, lower=None, upper=None, symbol="∑"):
    node = m_element("nary")
    props = m_element("naryPr")
    props.extend((m_element("chr", value=symbol), m_element("limLoc", value="undOvr")))
    node.append(props)
    lower_node = m_element("sub")
    if lower is not None:
        append_math(lower_node, lower)
    upper_node = m_element("sup")
    if upper is not None:
        append_math(upper_node, upper)
    expr = m_element("e")
    append_math(expr, expression)
    node.extend((lower_node, upper_node, expr))
    return node


def m_radical(expression):
    node = m_element("rad")
    props = m_element("radPr")
    props.append(m_element("degHide", value="1"))
    degree = m_element("deg")
    expr = m_element("e")
    append_math(expr, expression)
    node.extend((props, degree, expr))
    return node


def m_accent(expression, symbol="^"):
    node = m_element("acc")
    props = m_element("accPr")
    props.append(m_element("chr", value=symbol))
    expr = m_element("e")
    append_math(expr, expression)
    node.extend((props, expr))
    return node


def m_function(name: str, argument):
    node = m_element("func")
    func_name = m_element("fName")
    func_name.append(m_run(name, upright=True))
    expr = m_element("e")
    append_math(expr, m_delimiter(argument))
    node.extend((func_name, expr))
    return node


def m_o_math(expression):
    node = m_element("oMath")
    if expression.tag == qn("m:e"):
        for child in expression:
            node.append(deepcopy(child))
    else:
        node.append(deepcopy(expression))
    return node


def var(name: str):
    return m_run(name)


def upright(text: str):
    return m_run(text, upright=True)


def sub(base: str, index: str):
    return m_sub(var(base), var(index))


def sup(base: str, power: str):
    return m_sup(var(base), var(power))


def subsup(base: str, index: str, power: str):
    return m_subsup(var(base), var(index), var(power))


def build_inline_math(math_id: str):
    simple = {
        "c_o": sub("c", "o"),
        "c_max": sub("c", "max"),
        "c_hat_o": m_sub(m_accent(var("c")), var("o")),
        "l_lat": sub("L", "lat"),
        "l_ce": sub("L", "CE"),
        "lambda": var("λ"),
        "alpha": var("α"),
        "temperature": var("T"),
        "t_squared": sup("T", "2"),
        "pi_edge": m_subsup(var("π"), var("o"), upright("(i,j)")),
        "x": var("x"),
        "y": var("y"),
        "z_s": sub("z", "s"),
        "z_t": sub("z", "t"),
        "p_s_t": subsup("p", "s", "T"),
        "p_t_t": subsup("p", "t", "T"),
        "n_correct": sub("N", "correct"),
        "n_test": sub("N", "test"),
        "n_error": sub("N", "error"),
        "a": var("A"),
        "n": var("n"),
        "o": var("o"),
        "edge_ij": upright("(i,j)"),
    }
    factories = {
        "c_hat_definition": lambda: m_sequence(m_sub(m_accent(var("c")), var("o")), upright(" = "), m_fraction(sub("c", "o"), sub("c", "max"))),
        "c_max_positive": lambda: m_sequence(sub("c", "max"), upright(" > 0")),
        "lambda_set": lambda: m_sequence(var("λ"), upright(" ∈ {0.00, 0.05, 0.10, 0.20}")),
        "lambda_005": lambda: m_sequence(var("λ"), upright(" = 0.05")),
        "t_20": lambda: m_sequence(var("T"), upright(" = 20")),
        "alpha_05": lambda: m_sequence(var("α"), upright(" = 0.5")),
        "n_3": lambda: m_sequence(var("n"), upright(" = 3")),
        "n_test_834": lambda: m_sequence(sub("N", "test"), upright(" = 834")),
        "test_resolution": lambda: m_sequence(m_fraction(upright("100"), upright("834")), upright(" = 0.1199")),
        "one_e_minus_four": lambda: m_sequence(upright("1 × "), m_sup(upright("10"), upright("−4"))),
        "one_e_minus_three": lambda: m_sequence(upright("1 × "), m_sup(upright("10"), upright("−3"))),
        "three_e_minus_four": lambda: m_sequence(upright("3 × "), m_sup(upright("10"), upright("−4"))),
        "six_e_minus_four": lambda: m_sequence(upright("6 × "), m_sup(upright("10"), upright("−4"))),
        "one_e_minus_six": lambda: m_sequence(upright("1 × "), m_sup(upright("10"), upright("−6"))),
        "atol_1e4": lambda: m_sequence(var("atol"), upright(" = "), build_inline_math("one_e_minus_four")),
        "rtol_1e3": lambda: m_sequence(var("rtol"), upright(" = "), build_inline_math("one_e_minus_three")),
    }
    if math_id in simple:
        return deepcopy(simple[math_id])
    if math_id in factories:
        return factories[math_id]()
    raise ValueError(f"Unknown inline math marker: {math_id}")


def build_display_math(eq_id: str):
    edge_index = upright("(i,j)")
    pi_edge = m_subsup(var("π"), var("o"), edge_index)
    alpha_edge = m_subsup(var("α"), var("o"), edge_index)
    exp_alpha = m_function("exp", alpha_edge)

    equations = {
        "darts_mixture": lambda: m_sequence(
            pi_edge,
            upright(" = "),
            m_fraction(
                exp_alpha,
                m_nary(
                    m_function("exp", m_subsup(var("α"), upright("o′"), edge_index)),
                    lower=upright("o′"),
                ),
            ),
        ),
        "lut_normalization": lambda: m_sequence(
            m_sub(m_accent(var("c")), var("o")),
            upright(" = "),
            m_fraction(sub("c", "o"), sub("c", "max")),
            upright(",   "),
            sub("c", "max"),
            upright(" > 0"),
        ),
        "cell_latency": lambda: m_sequence(
            m_subsup(var("L"), upright("lat"), upright("cell")),
            upright(" = "),
            m_sub(m_run("mean", upright=True), edge_index),
            upright(" "),
            m_delimiter(
                m_nary(
                    m_sequence(pi_edge, upright(" "), m_sub(m_accent(var("c")), var("o"))),
                    lower=var("o"),
                ),
                opening="[",
                closing="]",
            ),
        ),
        "combined_latency": lambda: m_sequence(
            sub("L", "lat"),
            upright(" = "),
            m_fraction(
                m_sequence(m_subsup(var("L"), upright("lat"), upright("normal")), upright(" + "), m_subsup(var("L"), upright("lat"), upright("reduction"))),
                upright("2"),
            ),
        ),
        "architecture_objective": lambda: m_sequence(sub("L", "arch"), upright(" = "), sub("L", "CE"), upright(" + "), var("λ"), upright(" "), sub("L", "lat")),
        "student_distribution": lambda: m_sequence(
            subsup("p", "s", "T"), upright(" = "), m_function("softmax", m_fraction(sub("z", "s"), var("T")))
        ),
        "teacher_distribution": lambda: m_sequence(
            subsup("p", "t", "T"), upright(" = "), m_function("softmax", m_fraction(sub("z", "t"), var("T")))
        ),
        "kd_loss": lambda: m_sequence(
            sub("L", "KD"), upright(" = "), var("α"), upright(" "), sub("L", "CE"),
            m_delimiter(m_sequence(sub("z", "s"), upright(", "), var("y"))),
            upright(" + "), m_delimiter(m_sequence(upright("1 − "), var("α"))), upright(" "), sup("T", "2"), upright(" "),
            upright("KL"), m_delimiter(m_sequence(subsup("p", "t", "T"), upright(" ∥ "), subsup("p", "s", "T"))),
        ),
        "top1_accuracy": lambda: m_sequence(var("A"), upright(" = "), m_fraction(sub("N", "correct"), sub("N", "test"))),
        "sample_sd": lambda: m_sequence(
            var("s"), upright(" = "),
            m_radical(
                m_fraction(
                    m_nary(
                        m_sup(m_delimiter(m_sequence(m_sub(var("A"), var("i")), upright(" − "), m_accent(var("A"), symbol="¯"))), upright("2")),
                        lower=upright("i = 1"), upper=var("n"),
                    ),
                    m_sequence(var("n"), upright(" − 1")),
                )
            ),
            upright(",   "), var("n"), upright(" = 3"),
        ),
    }
    try:
        return equations[eq_id]()
    except KeyError as exc:
        raise ValueError(f"Unknown display equation marker: {eq_id}") from exc


DISPLAY_EQUATION_ORDER = [
    "darts_mixture",
    "lut_normalization",
    "cell_latency",
    "combined_latency",
    "architecture_objective",
    "student_distribution",
    "teacher_distribution",
    "kd_loss",
    "top1_accuracy",
    "sample_sd",
]


# Narrative-proposal base with the named ``ijcce_simple_manuscript`` override:
# single column; Times New Roman; left-aligned body; black hierarchy; no table
# shading or vertical rules. The first-page pattern is a restrained
# memo_masthead without metadata rows or a decorative rule.
SKIP_PREFIXES = (
    "Chen, Hsia, and Chen proposed",
    "Luo et al.'s AMPVNet",
    "Table 1 separates scientific role",
    "Hardware-aware NAS methods demonstrate",
    "An operator LUT offers",
    "The controlled KD layer",
    "The deployment stack uses",
    "This study is positioned around",
    "The decision task is",
    "Two evidence layers are maintained",
    "The manifest assigns eight images",
    "Static INT8 calibration uses",
    "The edge output is",
    "Network weights are optimized",
    "The KD protocol is",
    "The teacher remains",
    "The baseline students for KD",
    "Accuracy is evaluated independently",
    "The report lists all seed-level",
    "The reproducibility package retains",
)


def set_run_font(run, size: float, *, bold=None, italic=None, color="000000"):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_style(style, size, *, bold=False, italic=False, before=0, after=6, line=1.15):
    style.font.name = "Times New Roman"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = bold


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))
    set_run_font(run, 9, color="666666")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    values = {
        "top": ("single", "8", "000000"),
        "bottom": ("single", "8", "000000"),
        "insideH": ("single", "4", "BFBFBF"),
        "left": ("nil", "0", "FFFFFF"),
        "right": ("nil", "0", "FFFFFF"),
        "insideV": ("nil", "0", "FFFFFF"),
    }
    for name, (value, size, color) in values.items():
        node = borders.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            borders.append(node)
        node.set(qn("w:val"), value)
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_header_bottom_border(row):
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), "000000")
        borders.append(bottom)
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def add_inline_markdown(paragraph, text, *, size=11, color="000000"):
    token_re = re.compile(r"(\[\[MATH:[a-z0-9_]+\]\]|\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")
    cursor = 0
    for match in token_re.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size, color=color)
        token = match.group(0)
        if token.startswith("[[MATH:"):
            math_id = MATH_MARKER_RE.fullmatch(token).group(1)
            paragraph._p.append(m_o_math(build_inline_math(math_id)))
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size, bold=True, color=color)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size, italic=True, color=color)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size, color=color)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size, color=color)


def add_display_equation(doc, eq_id: str, number: int):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)

    lead = paragraph.add_run("\t")
    set_run_font(lead, 10.5)

    math_para = m_element("oMathPara")
    math_para_pr = m_element("oMathParaPr")
    math_para_pr.append(m_element("jc", value="centerGroup"))
    math_para.append(math_para_pr)
    math_para.append(m_o_math(build_display_math(eq_id)))
    paragraph._p.append(math_para)

    number_run = paragraph.add_run(f"\t({number})")
    set_run_font(number_run, 10.5)


def parse_table(block):
    rows = []
    for line in block.splitlines():
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [item.strip() for item in line.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def table_widths(index, columns):
    presets = {
        1: [1450, 1450, 2050, 1800, 2610],
        2: [1850, 1350, 1050, 1050, 4060],
        3: [2350, 7010],
        4: [1600, 2860, 2450, 2450],
        5: [1650, 3650, 4060],
    }
    widths = presets.get(index)
    if widths and len(widths) == columns:
        return widths
    base = 9360 // columns
    return [base] * (columns - 1) + [9360 - base * (columns - 1)]


def add_table(doc, rows, index):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if (r_idx == 0 or re.fullmatch(r"[0-9.,/ x%-]+", value))
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            add_inline_markdown(paragraph, value, size=8.3 if len(rows[0]) >= 4 else 8.7)
            if r_idx == 0:
                for run in paragraph.runs:
                    run.bold = True
    set_table_geometry(table, table_widths(index, len(rows[0])))
    set_table_borders(table)
    set_header_bottom_border(table.rows[0])
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)
    return table


def compact_blocks(text):
    blocks = re.split(r"\n\s*\n", text)
    result = []
    for block in blocks:
        plain = " ".join(block.split())
        if not plain:
            continue
        if any(plain.startswith(prefix) for prefix in SKIP_PREFIXES):
            continue
        result.append(block.strip())
    return result


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False

    doc.core_properties.title = "Hardware-Aware Progressive Differentiable Architecture Search for Compact Palm-Vein Identification on Raspberry Pi"
    doc.core_properties.subject = "IJCCE working manuscript, Sections 1-3"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "palm-vein identification; hardware-aware NAS; P-DARTS; Raspberry Pi"

    configure_style(doc.styles["Normal"], 11, before=0, after=6, line=1.15)
    configure_style(doc.styles["Heading 1"], 13, bold=True, before=14, after=6, line=1.0)
    configure_style(doc.styles["Heading 2"], 11.5, bold=True, before=11, after=4, line=1.0)
    configure_style(doc.styles["Heading 3"], 11, bold=True, italic=True, before=9, after=3, line=1.0)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_field(footer, "PAGE")

    blocks = compact_blocks(SOURCE.read_text(encoding="utf-8"))
    source_text = SOURCE.read_text(encoding="utf-8")
    all_display_ids = DISPLAY_MARKER_RE.findall(source_text)
    if all_display_ids != DISPLAY_EQUATION_ORDER:
        raise ValueError(
            "Display equation markers must occur exactly once in the locked order: "
            f"{DISPLAY_EQUATION_ORDER}; found {all_display_ids}"
        )
    all_inline_markers = MATH_MARKER_RE.findall(source_text)
    for math_id in all_inline_markers:
        build_inline_math(math_id)

    table_index = 0
    equation_index = 0
    for idx, block in enumerate(blocks):
        if block.startswith("# "):
            heading = block[2:].strip()
            if idx == 0:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(12)
                paragraph.paragraph_format.keep_with_next = True
                run = paragraph.add_run(heading)
                set_run_font(run, 17.5, bold=True)
            else:
                paragraph = doc.add_paragraph(heading, style="Heading 1")
                paragraph.paragraph_format.page_break_before = heading.startswith(
                    ("Internal Figure Production Checklist", "References")
                )
        elif block.startswith("## "):
            doc.add_paragraph(block[3:].strip(), style="Heading 2")
        elif block.startswith("### "):
            doc.add_paragraph(block[4:].strip(), style="Heading 3")
        elif block.startswith("> "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Inches(0.22)
            paragraph.paragraph_format.right_indent = Inches(0.22)
            paragraph.paragraph_format.space_after = Pt(12)
            paragraph.paragraph_format.line_spacing = 1.0
            add_inline_markdown(paragraph, block[2:].strip(), size=9, color="555555")
            for run in paragraph.runs:
                run.italic = True
        elif block.startswith("|"):
            table_index += 1
            add_table(doc, parse_table(block), table_index)
        elif block.startswith("**Table "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.keep_with_next = True
            add_inline_markdown(paragraph, block, size=9.5)
        elif block.startswith("*Table note."):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.0
            add_inline_markdown(paragraph, block, size=8.5, color="555555")
        elif block.startswith("**Figure "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(7)
            paragraph.paragraph_format.keep_with_next = False
            paragraph.paragraph_format.line_spacing = 1.0
            add_inline_markdown(paragraph, block, size=9)
        elif block.startswith("*Figure production note"):
            # Production instructions belong in the removable internal checklist,
            # never in the manuscript body beside a figure caption.
            continue
        elif DISPLAY_MARKER_RE.fullmatch(block.strip()):
            equation_index += 1
            add_display_equation(doc, DISPLAY_MARKER_RE.fullmatch(block.strip()).group(1), equation_index)
        elif re.match(r"^\[\d+\] ", block):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.28)
            paragraph.paragraph_format.first_line_indent = Inches(-0.28)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.0
            add_inline_markdown(paragraph, block, size=9)
        else:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.widow_control = True
            if idx + 1 < len(blocks) and DISPLAY_MARKER_RE.fullmatch(blocks[idx + 1].strip()):
                paragraph.paragraph_format.keep_with_next = True
            add_inline_markdown(paragraph, " ".join(block.splitlines()), size=11)

    if equation_index != len(DISPLAY_EQUATION_ORDER):
        raise ValueError(f"Expected 10 display equations, converted {equation_index}")
    remaining = [
        marker
        for paragraph in doc.paragraphs
        for marker in re.findall(r"\[\[(?:EQ|MATH):[^\]]+\]\]", paragraph.text)
    ]
    if remaining:
        raise ValueError(f"Unconverted equation markers remain: {remaining}")

    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
