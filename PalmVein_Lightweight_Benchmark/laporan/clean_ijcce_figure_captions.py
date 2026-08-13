"""Shorten IJCCE captions and remove inline production notes."""

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "IJCCE_Manuscript_Sections_1_3_Working_Draft_v4_Temporary_Thesis_Figures.docx"
OUTPUT = ROOT / "IJCCE_Manuscript_Sections_1_3_Working_Draft_v5_Concise_Figure_Captions.docx"


CAPTIONS = {
    1: (
        "Study workflow and data roles for hardware-aware architecture search, controlled training, "
        "knowledge distillation, quantization, and Raspberry Pi evaluation."
    ),
    2: (
        "Palm-vein image preparation: (a) region-of-interest extraction and (b) preprocessing to a "
        "224 x 224 model input."
    ),
    3: (
        "Raspberry Pi latency guidance: (a) latency lookup table construction and (b) integration into "
        "progressive differentiable architecture search."
    ),
    4: (
        "ONNX model export, static INT8 post-training quantization, accuracy validation, and Raspberry Pi "
        "latency benchmarking."
    ),
}


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_run_font(run, *, bold=False, italic=False, color=None) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if color is not None:
        run.font.color.rgb = color


def main() -> None:
    document = Document(SOURCE)

    if "Figure Caption" not in document.styles:
        style = document.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(9)
        style.paragraph_format.space_before = Pt(2)
        style.paragraph_format.space_after = Pt(7)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = False

    found = set()
    production_notes = []
    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        if text.startswith("Figure production note"):
            production_notes.append(paragraph)
            continue

        for number, caption in CAPTIONS.items():
            if text.startswith(f"Figure {number}."):
                paragraph.clear()
                paragraph.style = document.styles["Figure Caption"]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                prefix = paragraph.add_run(f"Figure {number}. ")
                set_run_font(prefix, bold=True)
                body = paragraph.add_run(caption)
                set_run_font(body)
                found.add(number)
                break

    if found != set(CAPTIONS):
        raise ValueError(f"Expected captions {sorted(CAPTIONS)}, found {sorted(found)}")
    if len(production_notes) != 4:
        raise ValueError(f"Expected 4 inline production notes, found {len(production_notes)}")

    for paragraph in production_notes:
        remove_paragraph(paragraph)

    for paragraph in document.paragraphs:
        if paragraph.style.name == "Figure Image" and paragraph.text.startswith("TEMPORARY THESIS ARTWORK"):
            first_run = paragraph.runs[0]
            first_run.text = "Temporary thesis artwork - replace before submission"
            set_run_font(first_run, italic=True, color=RGBColor(127, 127, 127))

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
