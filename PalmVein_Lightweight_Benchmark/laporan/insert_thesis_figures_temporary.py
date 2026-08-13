"""Insert temporary thesis artwork into the synchronized IJCCE v6 draft."""

from pathlib import Path
from tempfile import TemporaryDirectory
from zipfile import ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "IJCCE_Manuscript_Sections_1_3_Working_Draft_v6_Base.docx"
THESIS = ROOT.parents[1] / "backup laporan tesisi" / "Draft BACKUP_Mohammad_Taris_Syahir_Zul_Fahmi_6025242008_Final.docx"
OUTPUT = ROOT / "IJCCE_Manuscript_Sections_1_3_Working_Draft_v6_Split_LUT_Figures.docx"


FIGURES = {
    1: {
        "assets": [("word/media/image12.png", 4.50)],
        "alt": "Temporary thesis research-design diagram used as Figure 1.",
    },
    2: {
        "assets": [
            ("word/media/image13.png", 6.20),
            ("word/media/image14.png", 6.20),
        ],
        "alt": "Temporary thesis ROI-extraction and preprocessing diagrams used as Figure 2.",
    },
    3: {
        "assets": [("word/media/image16.png", 5.80)],
        "alt": "Temporary thesis diagram of operator-shape benchmarking and Raspberry Pi latency-LUT construction used as Figure 3.",
    },
    4: {
        "assets": [("word/media/image22.png", 5.80)],
        "alt": "Temporary thesis diagram of latency-LUT integration into hardware-aware P-DARTS used as Figure 4.",
    },
    5: {
        "assets": [("word/media/image24.png", 6.20)],
        "alt": "Temporary thesis ONNX-export, PTQ, and deployment-benchmark diagram used as Figure 5.",
    },
}


def set_alt_text(inline_shape, text: str, title: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", text)
    doc_pr.set("title", title)


def insert_figure(caption, number: int, asset_paths, alt: str) -> None:
    paragraph = caption.insert_paragraph_before(style="Figure Image")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True

    note = paragraph.add_run("Temporary thesis artwork - replace before submission")
    note.italic = True
    note.font.size = Pt(8)
    note.font.color.rgb = RGBColor(127, 127, 127)
    note.add_break()

    for index, (asset_path, width) in enumerate(asset_paths):
        if index:
            paragraph.add_run().add_break()
            spacer = paragraph.add_run()
            spacer.add_break()
        shape = paragraph.add_run().add_picture(str(asset_path), width=Inches(width))
        set_alt_text(shape, alt, f"Temporary Figure {number}")


def main() -> None:
    document = Document(SOURCE)

    if "Figure Image" not in document.styles:
        style = document.styles.add_style("Figure Image", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["Normal"]
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.0

    captions = {}
    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        for number in FIGURES:
            if text.startswith(f"Figure {number}."):
                captions[number] = paragraph
                break

    if set(captions) != set(FIGURES):
        raise ValueError(f"Expected captions {sorted(FIGURES)}, found {sorted(captions)}")

    with TemporaryDirectory(prefix="ijcce_thesis_figures_") as temp_dir:
        temp_dir = Path(temp_dir)
        extracted = {}
        with ZipFile(THESIS) as archive:
            for spec in FIGURES.values():
                for member, _ in spec["assets"]:
                    if member not in extracted:
                        output = temp_dir / Path(member).name
                        output.write_bytes(archive.read(member))
                        extracted[member] = output

        for number in sorted(FIGURES):
            spec = FIGURES[number]
            assets = [(extracted[name], width) for name, width in spec["assets"]]
            insert_figure(captions[number], number, assets, spec["alt"])

        document.save(OUTPUT)

    print(OUTPUT)


if __name__ == "__main__":
    main()
