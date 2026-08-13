"""Add explicit figure-production placeholders to the IJCCE working draft.

This script keeps the manuscript text and equations intact. It creates a new
DOCX and adds working-only production instructions immediately before each
planned figure caption, followed by a removable production checklist before
the references.
"""

from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "IJCCE_Manuscript_Sections_1_3_Working_Draft_v2.docx"
OUTPUT = ROOT / "IJCCE_Manuscript_Sections_1_3_Working_Draft_v3_Figure_Placeholders.docx"


FIGURES = {
    "Figure 1.": {
        "status": "REDRAW REQUIRED",
        "source": "Thesis Figure 3.1 (research design)",
        "action": (
            "Use the thesis figure only as source material. Redraw it for the article with the "
            "6,672/834/834 split, the 3,336/3,336 internal NAS split, separate scratch/pretrained/KD "
            "branches, training-only calibration, and distinct FP32/INT8 deployment paths."
        ),
    },
    "Figure 2.": {
        "status": "MERGE AND UPDATE",
        "source": "Thesis Figures 3.2 and 3.3 (ROI extraction and preprocessing)",
        "action": (
            "Combine author-owned intermediate palm images into one horizontal pipeline. Update the "
            "labels to Gaussian 7 x 7, Otsu mask, 15 x 15 morphology, 384 x 384 crop, CLAHE 2.0/8 x 8, "
            "min-max normalization, and Lanczos resize to 224 x 224."
        ),
    },
    "Figure 3.": {
        "status": "MERGE AND REDRAW",
        "source": "Thesis Figures 3.5 and 3.8 (latency LUT and LUT-integrated P-DARTS)",
        "action": (
            "Create one compact method diagram showing operator-shape probes, QDQ conversion, Raspberry Pi "
            "measurement, corrected aggregation, normalization, and the differentiable path into architecture "
            "parameters. Keep LUT guidance visually separate from final-model benchmarking."
        ),
    },
    "Figure 4.": {
        "status": "UPDATE AND REDRAW",
        "source": "Thesis Figure 3.10 (ONNX export and PTQ deployment)",
        "action": (
            "Add minimum-validation-loss checkpoint selection, ONNX checker/parity gate, the 834-image "
            "training-only calibration manifest, FP32 and INT8 accuracy paths, file hashing, and Raspberry Pi "
            "mean/median/P95 latency reporting."
        ),
    },
}


def set_cell_like_border(paragraph) -> None:
    """Give a paragraph a restrained working-note box without using a table."""
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F2F2")
    p_pr.append(shading)

    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "A6A6A6")
        borders.append(border)
    p_pr.append(borders)


def add_placeholder(caption, number: int, spec: dict[str, str]) -> None:
    paragraph = caption.insert_paragraph_before(style="Figure Placeholder")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_cell_like_border(paragraph)

    lead = paragraph.add_run(f"FIGURE {number} PLACEHOLDER - {spec['status']}\n")
    lead.bold = True
    lead.font.color.rgb = RGBColor(31, 78, 121)
    paragraph.add_run(f"Source material: {spec['source']}.\n").bold = True
    paragraph.add_run(f"Production action: {spec['action']}")


def insert_note_before(
    reference,
    text: str,
    *,
    style: Optional[str] = None,
    bold_prefix: Optional[str] = None,
):
    paragraph = reference.insert_paragraph_before(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        paragraph.add_run(bold_prefix).bold = True
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)
    return paragraph


def main() -> None:
    document = Document(SOURCE)

    if "Figure Placeholder" not in document.styles:
        style = document.styles.add_style("Figure Placeholder", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(9)
        style.font.color.rgb = RGBColor(64, 64, 64)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.left_indent = Pt(8)
        style.paragraph_format.right_indent = Pt(8)
        style.paragraph_format.line_spacing = 1.0

    found = set()
    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        for prefix, spec in FIGURES.items():
            if text.startswith(prefix):
                add_placeholder(paragraph, int(prefix.split()[1].rstrip(".")), spec)
                found.add(prefix)

    missing = set(FIGURES) - found
    if missing:
        raise ValueError(f"Figure captions not found: {sorted(missing)}")

    references = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.style.name == "Heading 1" and paragraph.text.strip() == "References"
    )

    heading = references.insert_paragraph_before(style="Heading 1")
    heading.add_run().add_break(WD_BREAK.PAGE)
    heading.add_run("Internal Figure Production Checklist - Remove Before Submission")

    insert_note_before(
        references,
        "Direct transfer unchanged: none. The thesis drawings are useful source material, but all four method "
        "figures require consolidation or correction before journal submission.",
        bold_prefix="Direct transfer unchanged:",
    )

    insert_note_before(references, "Required for Sections 1-3", style="Heading 2")
    required = [
        "Figure 1 - redraw thesis Figure 3.1 as the complete study framework and data-role diagram.",
        "Figure 2 - merge and update thesis Figures 3.2 and 3.3 using author-owned palm-image intermediates.",
        "Figure 3 - merge and redraw thesis Figures 3.5 and 3.8 as one LUT-to-search integration diagram.",
        "Figure 4 - update and redraw thesis Figure 3.10 as the export, PTQ, parity, and device-benchmark pipeline.",
    ]
    for item in required:
        insert_note_before(references, item, style="List Bullet")

    insert_note_before(references, "Generate Later for Results and Discussion", style="Heading 2")
    later = [
        "Final frozen architecture/genotype: generate from the selected configuration file after the final model is frozen; place in Section 4, not Methods.",
        "Accuracy-size-latency Pareto plot: generate reproducibly from the final three-seed FP32/INT8 result ledger; label scratch, pretrained, and KD protocols separately.",
        "PTQ impact plot: generate paired FP32-to-INT8 changes in accuracy, model size, and Raspberry Pi latency after every final deployment artifact passes validation.",
        "Optional graphical abstract: create separately only after the complete manuscript is stable; do not reuse a general-purpose generative-AI image.",
    ]
    for item in later:
        insert_note_before(references, item, style="List Bullet")

    insert_note_before(references, "Do Not Transfer to the Condensed Article", style="Heading 2")
    excluded = [
        "Thesis Figures 2.1-2.3: generic acquisition/NIR/CNN background; they add little methodological evidence.",
        "Thesis Figures 2.5-2.6 and 3.4, 3.6, 3.7, and 3.9: generic KD, PTQ, training, DARTS, and P-DARTS diagrams; their essential logic is already covered by the four consolidated figures, equations, and tables.",
        "Thesis Chapter 4 charts: do not copy until their numbers are reconciled with the final three-seed and Raspberry Pi artifact ledger.",
    ]
    for item in excluded:
        insert_note_before(references, item, style="List Bullet")

    insert_note_before(
        references,
        "Artwork requirements: export each final figure as a separate file; keep text outside images as editable captions; "
        "use uniform lettering; target at least 1000 dpi for line drawings or 500 dpi for mixed image-line artwork. "
        "Retain raw images, plotting scripts, and transformation history.",
        bold_prefix="Artwork requirements:",
    )

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
