#!/usr/bin/env python3
"""Deterministic structural checks for the IJCCE Sections 1-3 v6 draft."""

from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree


HERE = Path(__file__).resolve().parent
DOCX = HERE / "IJCCE_Manuscript_Sections_1_3_Working_Draft_v6_Split_LUT_Figures.docx"

NS = {
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}


def require(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def all_text(document: Document) -> str:
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(chunks)


def paragraph_index(paragraphs: list[str], prefix: str) -> int:
    matches = [index for index, text in enumerate(paragraphs) if text.startswith(prefix)]
    require(len(matches) == 1, f"Expected one paragraph beginning {prefix!r}, found {len(matches)}")
    return matches[0]


def main() -> None:
    document = Document(DOCX)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    text = all_text(document)

    captions = [
        text
        for text in paragraphs
        if re.match(r"^Figure [1-5]\. ", text)
    ]
    expected_captions = [
        "Figure 1. Study workflow and data roles for hardware-aware architecture search, controlled training, knowledge distillation, quantization, and Raspberry Pi evaluation.",
        "Figure 2. Palm-vein image preparation: (a) region-of-interest extraction and (b) preprocessing to a 224 x 224 model input.",
        "Figure 3. Construction of the Raspberry Pi latency lookup table from operator-shape benchmarks.",
        "Figure 4. Integration of the device-specific latency lookup table into the hardware-aware P-DARTS search objective.",
        "Figure 5. ONNX model export, static INT8 post-training quantization, accuracy validation, and Raspberry Pi latency benchmarking.",
    ]
    require(captions == expected_captions, f"Caption sequence mismatch: {captions}")
    require("Figure production note" not in text, "Inline figure-production note remains")
    require("(a) latency lookup table construction and (b) integration" not in text, "Combined LUT caption remains")

    section_34 = paragraph_index(paragraphs, "3.4. Target-Device Latency Characterization")
    section_352 = paragraph_index(paragraphs, "3.5.2. Latency-Regularized Search Objective")
    section_353 = paragraph_index(paragraphs, "3.5.3. Progressive Search Schedule")
    section_38 = paragraph_index(paragraphs, "3.8. ONNX Export and INT8 Quantization")
    section_39 = paragraph_index(paragraphs, "3.9. Evaluation and Statistical Reporting")
    figure_3 = paragraph_index(paragraphs, "Figure 3.")
    figure_4 = paragraph_index(paragraphs, "Figure 4.")
    figure_5 = paragraph_index(paragraphs, "Figure 5.")
    require(section_34 < figure_3 < section_352, "Figure 3 is not confined to Section 3.4")
    require(section_352 < figure_4 < section_353, "Figure 4 is not in Section 3.5.2")
    require(section_38 < figure_5 < section_39, "Figure 5 is not in Section 3.8")

    require("Probes are exported with ONNX opset 13" in text, "ONNX/QDQ probe paragraph is missing")
    for number in range(1, 6):
        caption_index = paragraph_index(paragraphs, f"Figure {number}.")
        prior_text = " ".join(paragraphs[max(0, caption_index - 8) : caption_index])
        require(f"Figure {number}" in prior_text, f"Figure {number} lacks a nearby textual callout")

    require(len(document.tables) == 5, f"Expected 5 tables, found {len(document.tables)}")
    require(len(document.inline_shapes) == 6, f"Expected 6 inline images, found {len(document.inline_shapes)}")
    require("[[EQ:" not in text and "[[MATH:" not in text, "Unconverted equation marker remains")
    require("Internal Figure Production Checklist - Remove Before Submission" in text, "Internal checklist is missing")

    with ZipFile(DOCX) as archive:
        xml = etree.fromstring(archive.read("word/document.xml"))
        math_paragraphs = xml.xpath(".//m:oMathPara", namespaces=NS)
        math_objects = xml.xpath(".//m:oMath", namespaces=NS)
        anchors = xml.xpath(".//wp:anchor", namespaces=NS)
        inlines = xml.xpath(".//wp:inline", namespaces=NS)
        doc_properties = xml.xpath(".//wp:docPr", namespaces=NS)

    require(len(math_paragraphs) == 10, f"Expected 10 display OMML paragraphs, found {len(math_paragraphs)}")
    require(len(math_objects) >= 10, f"Expected editable OMML objects, found {len(math_objects)}")
    require(len(anchors) == 0, f"Expected no floating anchors, found {len(anchors)}")
    require(len(inlines) == 6, f"Expected 6 inline drawings, found {len(inlines)}")
    require(
        all(node.get("descr") and node.get("title") for node in doc_properties),
        "Every image must have title and description alt text",
    )

    equation_numbers = [
        match.group(1)
        for paragraph in paragraphs
        if (match := re.fullmatch(r"\((10|[1-9])\)", paragraph))
    ]
    for number in range(1, 11):
        require(
            equation_numbers.count(str(number)) == 1,
            f"Equation number ({number}) does not occur exactly once",
        )

    references_start = paragraph_index(paragraphs, "References (temporary numbered drafting style)")
    body_text = "\n".join(paragraphs[:references_start])
    reference_text = "\n".join(paragraphs[references_start + 1 :])
    citation_set = {int(value) for value in re.findall(r"\[(\d+)\]", body_text)}
    reference_numbers = [int(value) for value in re.findall(r"^\[(\d+)\]", reference_text, flags=re.MULTILINE)]
    expected = set(range(1, 36))
    require(citation_set == expected, f"Citation set mismatch: {sorted(citation_set)}")
    require(reference_numbers == list(range(1, 36)), f"Reference sequence mismatch: {reference_numbers}")

    print("PASS")
    print(f"DOCX: {DOCX}")
    print("Figures: 5 captions, 6 inline images, 0 floating anchors")
    print(f"Equations: {len(math_paragraphs)} display OMML paragraphs, {len(math_objects)} total OMML objects")
    print("Tables: 5")
    print("References: 35; citation/reference sets match")


if __name__ == "__main__":
    main()
